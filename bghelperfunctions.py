import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import datetime
import time
import math
import scipy.ndimage.measurements as im_meas
from mpl_toolkits.axes_grid1 import make_axes_locatable
from docx import Document
from docx.shared import Inches

_interp_s = 30;

class TimeBandTargets(object):
    """object to define the times and BG values for target readings, and to store the percentage times within targets"""
    def __init__(self, 
                 time_band_name='Overnight', 
                 time_start_end=(0,7),
                 target_bg=(5.0,9.0)):
        self.time_band_name = time_band_name
        self.time_start_end = time_start_end
        self.target_bg = target_bg
        self.percentage_in_target = 100
        self.percentage_above_target = 0
        self.percentage_below_target = 0
        
    def set_percentage_above_target(self, val):
        if (val + self.percentage_in_target + self.percentage_below_target) > 100:
            self.percentage_above_target = 100 - self.percentage_in_target - self.percentage_below_target
        else:
            self.percentage_above_target = val;
            
    def set_percentage_in_target(self, val):
        if (val + self.percentage_above_target + self.percentage_below_target) > 100:
            self.percentage_in_target = 100 - self.percentage_above_target - self.percentage_below_target
        else:
            self.percentage_above_target = val;
            
    def set_percentage_below_target(self, val):
        if (val + self.percentage_in_target + self.percentage_above_target) > 100:
            self.percentage_below_target = 100 - self.percentage_in_target - self.percentage_above_target
        else:
            self.percentage_below_target = val;    
        
def minus_time(t1, t2):
    a = datetime.timedelta(hours=t1.hour, minutes=t1.minute, seconds=t1.second)
    b = datetime.timedelta(hours=t2.hour, minutes=t2.minute, seconds=t2.second)
    return a - b

def plus_time(t1, t2):
    a = datetime.timedelta(hours=t1.hour, minutes=t1.minute, seconds=t1.second)
    b = datetime.timedelta(hours=t2.hour, minutes=t2.minute, seconds=t2.second)
    return a + b
    
def mealtime(time_in, variability_in_mins):
    """model times for comparisson of information with/without CGM"""
    mu = datetime.timedelta(hours=time_in.hour, minutes=time_in.minute).total_seconds()
    sigma = variability_in_mins * 60 / 5
    meal_time_s = np.random.normal(mu, sigma)
    h = int(math.floor(meal_time_s / (60 * 60)))
    m = int(math.floor((meal_time_s - h * 60 * 60)/60))
    s = int(math.floor(meal_time_s - h * 60 * 60 - m * 60))
    if (h > 23):
        h = 23;
        m = 59;
        s = 59;
    meal_time = datetime.time(hour=h, minute=m, second=s)
    return meal_time
    
def add_daily_scatter(df, date_to_plot, ax):
    """for approximate comparison of information with/without CGM"""
    if 'Finger BG, mmoll-1' in df.columns:
        finger_df = df[(df['date']==date_to_plot) & (df['Finger BG, mmoll-1'].notnull())]
        ax.scatter(finger_df['time'].tolist(), finger_df['Finger BG, mmoll-1'].tolist())
    else:
        day_df = df.loc[df['date'] == date_to_plot]
        meal_times = np.asarray([datetime.time(hour=8, minute=0), 
					  datetime.time(hour=12, minute=30), 
					  datetime.time(hour=18, minute=30), 
					  datetime.time(hour=23, minute=0)])
        meal_variability_mins = np.asarray([30, 60, 90, 120])
        latest_available_time = day_df['time'].iloc[0]
        meal_variability_mins = meal_variability_mins[meal_times < latest_available_time].tolist()
        meal_times = meal_times[meal_times < latest_available_time].tolist()

        ts = [mealtime(t, r) for (t, r) in zip(meal_times, meal_variability_mins)]
        closest_ts = [day_df['datetime'][((day_df['time'].apply(lambda x: minus_time(x,t))).abs().argsort()[:1])] for t in ts]
        mus = [day_df['BG, mmoll-1'].get(t).values[0] for t in closest_ts]
        readings = [np.around(np.random.normal(mu, 0.25), 1) for mu in mus]
        ax.scatter([t.index.time[0] for t in closest_ts], readings)
    

def plot_daily_BG(df, date_to_plot, ax):
    """generate a line plot of BG readings with bands showing (hardcoded) target range"""
    ax.axhspan(2.0, 3.5, alpha=0.25, color='r')
    ax.axhspan(12.0, 18.0, alpha=0.25, color='r')
    ax.axhspan(8.0, 12.0, alpha=0.25, color='orange')
    df.loc[df['date'] == date_to_plot].set_index(df.loc[df['date'] == date_to_plot]['time'])['BG, mmoll-1'].dropna().plot(ax=ax)
    temp = df.loc[df['date'] == date_to_plot].set_index(df.loc[df['date'] == date_to_plot]['time'])['BG, mmoll-1'];
    ax.set_ylabel('BG, mmoll-1')
    ax.set_xlabel('Time')
    ax.set_ylim((2.0, 18.0))
    ax.set_xlim((datetime.time(hour=0), datetime.time(hour=23, minute=59)))
    t = [datetime.time(hour=4*x) for x in range(6)]
    t.append(datetime.time(hour=23, minute=59))
    ax.set_xticks(t)
    #add_daily_scatter(df, date_to_plot,  ax)
    
def lastWday(adate, w):
    MON, TUE, WED, THU, FRI, SAT, SUN = range(7)
    """Mon:w=0, Sun:w=6"""
    delta = (adate.weekday() + 6 - w) % 7 + 1
    return adate - datetime.timedelta(days=delta)
       
def roundTimeToLastXMinutes(tm, X):
    tm = datetime.datetime.combine(datetime.date.today(), tm)
    tm = tm - datetime.timedelta(minutes=tm.minute % X,
                             seconds=tm.second,
                             microseconds=tm.microsecond)
    tm = datetime.time(hour=tm.hour, minute=tm.minute)
    return tm
	
def plot_hypos(df, startdate, enddate=datetime.date.today(), hypo_max_bg=3.5, outpath=None, desc_str=None, document=None):
	"""show colourmap and provide basis for extracting stats on separate hypoglycaemic episodes"""
	sample_df = df.loc[(df['date'] >= startdate) & (df['date'] <= enddate)];

	# ensure even sampling
	t = sample_df.index
	r = pd.date_range(t.min(), t.max(), freq='{}S'.format(_interp_s))
	interp_df = sample_df.reindex(r).interpolate('index')

	# binarise hypos and label each distinct hypo event
	mask = (interp_df['BG, mmoll-1'] < hypo_max_bg).values
	lbl, nfeat = im_meas.label((mask).astype(int))
	o = im_meas.find_objects(lbl)
	hypo_starts_t = np.asarray([interp_df.index[x[0].start].time() for x in o])
	hypo_starts_dt = [interp_df.index[x[0].start] for x in o]
	#durations = [minus_time(interp_df.index[x[0].stop].time(), interp_df.index[x[0].start].time()) for x in o]
	#print([d.total_seconds() for d in durations])

	days = np.unique(interp_df.index.date)
	zero = datetime.datetime(1990, 1, 1, 0, 0, 0).time();
	start_index = math.floor(minus_time(interp_df.iloc[0]['time'], zero).total_seconds()/30);
	hypo_arr = 10 * np.ones((len(days) * (2*60*24)))
	vals = interp_df['BG, mmoll-1'].values;
	hypo_arr[start_index:start_index+len(interp_df['BG, mmoll-1'].values)] = vals;
	hypo_arr = hypo_arr.reshape((len(days), (2*60*24)));

	hypo_arr = np.ma.masked_where(hypo_arr>hypo_max_bg, hypo_arr)

	fig, ax = plt.subplots(1,1)
	cbdum = ax.imshow(hypo_arr, aspect='auto', cmap='Reds_r', clim=(2.0,3.5), interpolation=None)
	fig.subplots_adjust(top=0.9, right=0.9, hspace=0.1);
	cax = fig.add_axes([0.95, ax.get_position().y0, 0.03, ax.get_position().y1 - ax.get_position().y0])
	fig.colorbar(cbdum, cax=cax);

	t = [x * 4 * 2 * 60 for x in range(6)]
	d = [x * 4 for x in range(int(len(days)/4))]
	d.append(hypo_arr.shape[0]-1)
	ax.set_xticks(t)
	ax.set_yticks(d)
	tlbl = [datetime.time(hour= int(math.floor(tx / (2 * 60))), minute=0) for tx in t]
	dlbl = [x.strftime("%Y-%b-%d") for x in days[d]]
	ax.set_xticklabels(tlbl)
	ax.set_xlabel("Time of day")
	ax.set_yticklabels(dlbl)
	ax.set_ylabel("Date")
	ax.set_title(startdate.strftime("Hypoglycaemic episodes in period %b %d, %Y - " + enddate.strftime("%b %d, %Y")))
	cax.set_ylabel("BG, mmoll-1")
	if outpath:
		plt.savefig(outpath + '/{}hypos.png'.format(desc_str), dpi=600)
		if document:
			document.add_page_break()
			document.add_heading("{}hypos".format(desc_str), level=2)
			document.add_picture(outpath + '/{}hypos.png'.format(desc_str), width=Inches(6.2))
	plt.show()

	return ax

def percentageTimeInTarget(df, startdate, enddate, time_band_target_list, outpath=None, desc_str=None, document=None):
    """plot % time below/in/above targets, using time-banded targets object"""
    sample_df = df.loc[(df['date'] >= startdate) & (df['date'] <= enddate)]
    
    for tidx, time_band_target in enumerate(time_band_target_list):
        tb_times = time_band_target.time_start_end
        tb_targets = time_band_target.target_bg
        band_start_t = datetime.time(hour=int(math.floor(tb_times[0])), minute=int(tb_times[0]-math.floor(tb_times[0])))
        band_end_t = datetime.time(hour=int(math.floor(tb_times[1])), minute=int(tb_times[1]-math.floor(tb_times[1])))
        sub_sample_df = sample_df.loc[(sample_df['time'] >= band_start_t) & (sample_df['time'] <= band_end_t)]['BG, mmoll-1']
        
        t = sub_sample_df.index
        r = pd.date_range(t.min(), t.max(), freq='{}S'.format(_interp_s))
        interp_df = sub_sample_df.reindex(t.union(r)).interpolate('index')
        
        time_band_target_list[tidx].percentage_below_target = np.round(100.0 * (interp_df < tb_targets[0]).sum() / interp_df.count(),1)
        time_band_target_list[tidx].percentage_in_target = np.round(100.0 * ((interp_df >= tb_targets[0]) & (interp_df <= tb_targets[1])).sum() / interp_df.count(),1)
        time_band_target_list[tidx].percentage_above_target = np.round(100.0 * (interp_df > tb_targets[1]).sum() / interp_df.count(),1)
        
    
    # handle plotting
    fig, ax = plt.subplots(1,1)
    idx = len(time_band_target_list) - np.arange(len(time_band_target_list));
    xindices = np.arange(len(time_band_target_list))
    bar_w = 0.35
    rects = []
    # note that repeat looping here isn't ideal, but such a small penalty it's not worth fixing
    above_lims = [tbt.percentage_above_target for tbt in time_band_target_list]
    in_lims = [tbt.percentage_in_target for tbt in time_band_target_list]
    below_lims = [tbt.percentage_below_target for tbt in time_band_target_list]
    
    rects.append(ax.bar(xindices, 
                        below_lims, 
                        bar_w, 
                        color='orange', 
                        label='Below target'))
    rects.append(ax.bar(xindices, 
                        in_lims, 
                        bar_w, 
                        color='g', 
                        bottom=below_lims, 
                        label='Within target'))
    rects.append(ax.bar(xindices, 
                        above_lims, 
                        bar_w, 
                        color='r', 
                        bottom=[sum(x) for x in zip(below_lims, in_lims)], 
                        label='Above target'))
    ax.set_xticks(xindices)
    ax.set_xticklabels([tbt.time_band_name for tbt in time_band_target_list])
    ax.set_ylabel("Percentage time in each range")
    ax.set_xlabel("Periods of interest")
    ax.legend()
    ax.set_title(startdate.strftime("Time spent within target ranges in period %b %d, %Y - " + enddate.strftime("%b %d, %Y")))
    
    # handle labeling
    labels = []
    for cidx, container in enumerate(rects):
        for ridx, rect in enumerate(container):
            tbt = time_band_target_list[ridx];
            pc_lbls = (tbt.percentage_below_target, tbt.percentage_in_target, tbt.percentage_above_target);

            if (rect.get_height() > 3):
                yloc = rect.get_y() + rect.get_height()/2.0
                xloc = rect.get_x() + rect.get_width()/2.0

                pc_str = str.format("{0}", pc_lbls[cidx])
                label = ax.text(xloc, yloc, pc_str, horizontalalignment='center',
                             verticalalignment='center', color='w', weight='bold',
                             clip_on=True)
                labels.append(label)
    if outpath:
        plt.savefig(outpath + '/{}targets.png'.format(desc_str), dpi=600)
        if document:
            document.add_page_break()
            document.add_heading("{}targets".format(desc_str), level=2)
            document.add_picture(outpath + '/{}targets.png'.format(desc_str), width=Inches(6.2))
    return time_band_target_list;
	
def rolling_averages(series_groupby, filt_type='median', filt_half_size_mins=60, deltat_seconds=30, q=0.5):
	if filt_type=='median':
		x = series_groupby.median();
	elif filt_type=='mean':
		x = series_groupby.mean();
	elif filt_type=='quantile':
		x = series_groupby.quantile(q);
	t = x.index;
	out = np.zeros((2 * int(filt_half_size_mins * (60/deltat_seconds))+1, len(x)));
	for idx in range(out.shape[0]):
		out[idx] = np.roll(x.values, idx - int(filt_half_size_mins * (60/deltat_seconds)));
	if filt_type=='median':
		return t, np.median(out, axis=0);
	elif filt_type=='quantile':
		return t, np.quantile(out, q, axis=0);
	elif filt_type=='mean':
		return t, np.mean(out, axis=0);

def plot_long_term_BG(df, startdate, enddate, time_smoothing_s=600, outpath=None, desc_str=None, document=None):
	"""plot (rolling) median BG between two dates, along with 25-75 percentile range"""
	sample_df = df.loc[(df['date'] > startdate) & (df['date'] < enddate)]
	if minus_time(sample_df.iloc[0]['time'], 
						sample_df.iloc[1]['time']).total_seconds() > time_smoothing_s:
		interp_df = sample_df.reindex(pd.date_range(sample_df.index.min(), 
													sample_df.index.max(), 
													freq='{}S'.format(time_smoothing_s))).interpolate('index');
	else:
		interp_df = sample_df.resample('{}S'.format(time_smoothing_s)).mean();
	interp_df['time'] = interp_df.index.time;
	groupbytime = interp_df.groupby('time')['BG, mmoll-1']
	fig, ax = plt.subplots(1,1)
	ax.axhspan(2.0, 3.5, alpha=0.25, color='r')
	ax.axhspan(12.0, 18.0, alpha=0.25, color='r')
	ax.axhspan(8.0, 12.0, alpha=0.25, color='orange')
	_, med = rolling_averages(groupbytime, 
								filt_half_size_mins=60, 
								filt_type='median', 
								deltat_seconds=time_smoothing_s);
	_, lq = rolling_averages(groupbytime, 
								filt_half_size_mins=60, 
								filt_type='quantile',
								deltat_seconds=time_smoothing_s, 
								q=0.25);
	t, uq = rolling_averages(groupbytime, 
								filt_half_size_mins=60, 
								filt_type='quantile', 
								deltat_seconds=time_smoothing_s, 
								q=0.75);
	_, q10 = rolling_averages(groupbytime, 
								filt_half_size_mins=60, 
								filt_type='quantile', 
								deltat_seconds=time_smoothing_s, 
								q=0.1);								
	_, q90 = rolling_averages(groupbytime, 
								filt_half_size_mins=60, 
								filt_type='quantile', 
								deltat_seconds=time_smoothing_s, 
								q=0.9);
	bgplt = ax.plot(t, med, label='Median');
	ax.fill_between(t, q10, q90, alpha=0.1, color='gray', label='10-90% quantile')
	ax.fill_between(t, lq, uq, alpha=0.25, color='gray', label='25-75% quantile')
	ax.legend()
	ax.set_ylabel('BG, mmoll-1')
	ax.set_xlabel('Time')
	ax.set_ylim((2.0, 18.0))
	ax.set_xlim((datetime.time(hour=0), datetime.time(hour=23, minute=59)))
	t = [datetime.time(hour=4*x) for x in range(6)]
	t.append(datetime.time(hour=23, minute=59))
	ax.set_xticks(t)
	ax.set_title(startdate.strftime("Median BG in period %b %d, %Y - " + enddate.strftime("%b %d, %Y")))
	if outpath:
		plt.savefig(outpath + '/{}quantiles.png'.format(desc_str), dpi=600)
		if document:
			document.add_page_break()
			document.add_heading("{}quantiles".format(desc_str), level=2)
			document.add_picture(outpath + '/{}quantiles.png'.format(desc_str), width=Inches(6.2))
	return ax;
